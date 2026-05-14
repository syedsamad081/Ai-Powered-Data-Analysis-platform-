"""
services/file_handler.py  —  Reads uploaded files into a pandas DataFrame.

This is the first step after a file is uploaded. Depending on the file type,
a different library is used to read it:

  .csv          → pandas read_csv  (handles UTF-8 and latin-1 encoding)
  .xlsx / .xls  → pandas read_excel with the openpyxl engine
  .docx         → python-docx  (extracts the first table found in the document)
  .pdf          → pdfplumber   (extracts the first table found across all pages)

All parsers return a standard pandas DataFrame so the rest of the app
doesn't need to know what file type was originally uploaded.
"""

import pandas as pd
import pdfplumber
from docx import Document


def parse_file(filepath: str) -> pd.DataFrame:
    """
    Read a file from disk and return it as a pandas DataFrame.
    The file type is determined from the file extension.

    Raises ValueError if the file type is not supported or no table is found.
    """
    ext = filepath.rsplit(".", 1)[-1].lower()

    if ext == "csv":
        return _parse_csv(filepath)
    if ext in ("xlsx", "xls"):
        return _parse_excel(filepath)
    if ext == "docx":
        return _parse_docx(filepath)
    if ext == "pdf":
        return _parse_pdf(filepath)

    raise ValueError(f"Unsupported file type: .{ext}")


# ── Individual parsers ────────────────────────────────────────

def _parse_csv(filepath: str) -> pd.DataFrame:
    """
    Read a CSV file.
    Tries UTF-8 first (most common), then falls back to latin-1
    which handles files with special characters like é, ü, ñ.
    """
    try:
        return pd.read_csv(filepath, encoding="utf-8")
    except UnicodeDecodeError:
        return pd.read_csv(filepath, encoding="latin-1")


def _parse_excel(filepath: str) -> pd.DataFrame:
    """
    Read an Excel file (.xlsx or .xls).
    Uses the openpyxl engine which supports modern .xlsx format.
    Reads the first sheet by default.
    """
    return pd.read_excel(filepath, engine="openpyxl")


def _parse_docx(filepath: str) -> pd.DataFrame:
    """
    Extract the first table from a Word document (.docx).
    The first row of the table is treated as the column headers,
    and all remaining rows become the data.

    Raises ValueError if the document contains no tables.
    """
    doc = Document(filepath)

    if not doc.tables:
        raise ValueError("No tables found in the DOCX file. "
                         "The document must contain at least one table.")

    # Get the first table in the document
    table = doc.tables[0]

    # Convert each row to a list of cell text values
    data = [[cell.text.strip() for cell in row.cells] for row in table.rows]

    if not data:
        raise ValueError("The table in the DOCX file is empty.")

    # First row = column names, rest = data rows
    df = pd.DataFrame(data[1:], columns=data[0])

    # Try to convert each column to a number — columns that can't be
    # converted (e.g. names, labels) stay as text
    return _coerce_numeric(df)


def _parse_pdf(filepath: str) -> pd.DataFrame:
    """
    Extract the first table found across all pages of a PDF.
    Uses pdfplumber which can detect table borders and text positions.

    Raises ValueError if no usable table is found in the PDF.
    """
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                # A valid table must have at least a header row and one data row
                if table and len(table) >= 2:
                    # Clean up cell text (strip whitespace, replace None with "")
                    cleaned = [[(cell or "").strip() for cell in row] for row in table]
                    headers = cleaned[0]
                    rows    = cleaned[1:]
                    df = pd.DataFrame(rows, columns=headers)
                    return _coerce_numeric(df)

    raise ValueError("No usable tables found in the PDF file. "
                     "Make sure the PDF contains a visible data table.")


def _coerce_numeric(df: pd.DataFrame) -> pd.DataFrame:
    """
    Try to convert every column to a number.
    Columns that contain text (like names or categories) will stay as strings.
    This is used after reading Word and PDF files, where all values start as strings.
    """
    for col in df.columns:
        try:
            df[col] = pd.to_numeric(df[col])
        except (ValueError, TypeError):
            pass    # column stays as text — that's fine
    return df
